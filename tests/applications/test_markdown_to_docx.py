from pathlib import Path

import pytest
from docx import Document

from applications.markdown_to_docx import (
    DEFAULT_TEMPLATE_ANCHOR,
    StyleProfile,
    TemplateCatalog,
    create_starter_template,
    render_markdown_to_docx,
)
from applications.markdown_to_docx.cli import main as cli_main


def _non_empty_paragraph_texts(doc: Document) -> list[str]:
    return [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]


def test_render_markdown_to_docx_preserves_template_and_renders_markdown(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("Cover")
    doc.add_paragraph(DEFAULT_TEMPLATE_ANCHOR)
    doc.add_paragraph("Appendix")
    doc.save(template_path)

    output_path = tmp_path / "output.docx"
    markdown_text = """# Report Title

This is the first paragraph.

| Item | Value |
| --- | --- |
| OCR | Ready |
"""

    render_markdown_to_docx(
        markdown_text=markdown_text,
        output_path=output_path,
        template_path=template_path,
    )

    rendered = Document(output_path)
    paragraph_texts = _non_empty_paragraph_texts(rendered)
    assert DEFAULT_TEMPLATE_ANCHOR not in paragraph_texts
    assert paragraph_texts == [
        "Cover",
        "Report Title",
        "This is the first paragraph.",
        "Appendix",
    ]
    assert len(rendered.tables) == 1
    assert rendered.tables[0].cell(0, 0).text == "Item"
    assert rendered.tables[0].cell(1, 0).text == "OCR"
    assert rendered.tables[0].cell(1, 1).text == "Ready"


def test_render_markdown_to_docx_applies_style_profile(tmp_path: Path):
    output_path = tmp_path / "styled.docx"
    profile = StyleProfile.from_dict(
        {
            "normal": {
                "font_name": "Arial",
                "font_size_pt": 11,
                "line_spacing": 1.4,
            },
            "heading1": {
                "font_name": "Arial",
                "font_size_pt": 20,
                "bold": True,
            },
        }
    )

    render_markdown_to_docx(
        markdown_text="# Styled Title\n\nStyled body paragraph.",
        output_path=output_path,
        style_profile=profile,
    )

    rendered = Document(output_path)
    title_paragraph = next(
        paragraph for paragraph in rendered.paragraphs if paragraph.text == "Styled Title"
    )
    body_paragraph = next(
        paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text == "Styled body paragraph."
    )

    assert title_paragraph.runs[0].font.bold is True
    assert round(title_paragraph.runs[0].font.size.pt) == 20
    assert round(body_paragraph.runs[0].font.size.pt) == 11
    assert body_paragraph.paragraph_format.line_spacing == pytest.approx(1.4)


def test_template_catalog_resolves_template_and_style_paths(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    create_starter_template(template_path, title="Catalog Template")
    style_path = tmp_path / "style.json"
    style_path.write_text(
        """
{
  "normal": {"font_name": "Calibri", "font_size_pt": 12},
  "heading1": {"font_name": "Calibri", "font_size_pt": 18, "bold": true}
}
""".strip(),
        encoding="utf-8",
    )
    catalog_path = tmp_path / "catalog.json"
    catalog_path.write_text(
        f"""
{{
  "templates": {{
    "report": {{
      "template_path": "{template_path.name}",
      "style_profile_path": "{style_path.name}",
      "anchor": "{DEFAULT_TEMPLATE_ANCHOR}"
    }}
  }}
}}
""".strip(),
        encoding="utf-8",
    )

    catalog = TemplateCatalog.load(catalog_path)
    resolved = catalog.resolve("report")

    assert resolved.template_path == template_path
    assert resolved.style_profile_path == style_path
    assert resolved.anchor == DEFAULT_TEMPLATE_ANCHOR


def test_cli_render_uses_default_anchor_for_template(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    create_starter_template(template_path, title="CLI Template")
    markdown_path = tmp_path / "input.md"
    markdown_path.write_text("# CLI Title\n\nCLI paragraph.", encoding="utf-8")
    output_path = tmp_path / "output.docx"

    exit_code = cli_main(
        [
            "render",
            "--markdown",
            str(markdown_path),
            "--output",
            str(output_path),
            "--template",
            str(template_path),
        ]
    )

    assert exit_code == 0
    rendered = Document(output_path)
    assert "CLI Title" in _non_empty_paragraph_texts(rendered)
