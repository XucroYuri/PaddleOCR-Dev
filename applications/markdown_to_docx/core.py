from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable, Optional

from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import markdown

from ppstructure.recovery.table_process import HtmlToDocx

DEFAULT_TEMPLATE_ANCHOR = "{{PADDLEOCR_MARKDOWN}}"


@dataclass(frozen=True)
class BlockStyle:
    font_name: str = "Calibri"
    east_asia_font: Optional[str] = None
    font_size_pt: float = 11.0
    bold: bool = False
    italic: bool = False
    underline: bool = False
    alignment: str = "left"
    line_spacing: Optional[float] = 1.25
    space_before_pt: float = 0.0
    space_after_pt: float = 6.0

    @classmethod
    def from_dict(cls, data: dict[str, Any] | None, /, **defaults: Any) -> "BlockStyle":
        if data is None:
            data = {}
        return cls(**{**defaults, **data})


@dataclass(frozen=True)
class StyleProfile:
    normal: BlockStyle = field(default_factory=BlockStyle)
    heading1: BlockStyle = field(
        default_factory=lambda: BlockStyle(font_size_pt=18.0, bold=True, space_before_pt=12.0, space_after_pt=8.0)
    )
    heading2: BlockStyle = field(
        default_factory=lambda: BlockStyle(font_size_pt=16.0, bold=True, space_before_pt=10.0, space_after_pt=6.0)
    )
    heading3: BlockStyle = field(
        default_factory=lambda: BlockStyle(font_size_pt=14.0, bold=True, space_before_pt=8.0, space_after_pt=4.0)
    )
    list_item: BlockStyle = field(default_factory=BlockStyle)
    code_block: BlockStyle = field(
        default_factory=lambda: BlockStyle(font_name="Courier New", font_size_pt=10.0, space_before_pt=6.0, space_after_pt=6.0)
    )
    quote: BlockStyle = field(
        default_factory=lambda: BlockStyle(italic=True, space_before_pt=6.0, space_after_pt=6.0)
    )
    table_header: BlockStyle = field(
        default_factory=lambda: BlockStyle(font_size_pt=11.0, bold=True)
    )
    table_cell: BlockStyle = field(default_factory=BlockStyle)

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "StyleProfile":
        return cls(
            normal=BlockStyle.from_dict(data.get("normal")),
            heading1=BlockStyle.from_dict(
                data.get("heading1"),
                font_size_pt=18.0,
                bold=True,
                space_before_pt=12.0,
                space_after_pt=8.0,
            ),
            heading2=BlockStyle.from_dict(
                data.get("heading2"),
                font_size_pt=16.0,
                bold=True,
                space_before_pt=10.0,
                space_after_pt=6.0,
            ),
            heading3=BlockStyle.from_dict(
                data.get("heading3"),
                font_size_pt=14.0,
                bold=True,
                space_before_pt=8.0,
                space_after_pt=4.0,
            ),
            list_item=BlockStyle.from_dict(data.get("list_item")),
            code_block=BlockStyle.from_dict(
                data.get("code_block"),
                font_name="Courier New",
                font_size_pt=10.0,
                space_before_pt=6.0,
                space_after_pt=6.0,
            ),
            quote=BlockStyle.from_dict(
                data.get("quote"),
                italic=True,
                space_before_pt=6.0,
                space_after_pt=6.0,
            ),
            table_header=BlockStyle.from_dict(
                data.get("table_header"),
                font_size_pt=11.0,
                bold=True,
            ),
            table_cell=BlockStyle.from_dict(data.get("table_cell")),
        )

    @classmethod
    def load(cls, path: str | Path) -> "StyleProfile":
        with Path(path).open("r", encoding="utf-8") as fp:
            return cls.from_dict(json.load(fp))


@dataclass(frozen=True)
class ResolvedTemplate:
    template_path: Path
    style_profile_path: Optional[Path] = None
    anchor: str = DEFAULT_TEMPLATE_ANCHOR


@dataclass(frozen=True)
class TemplateCatalog:
    root: Path
    templates: dict[str, ResolvedTemplate]

    @classmethod
    def load(cls, path: str | Path) -> "TemplateCatalog":
        catalog_path = Path(path).resolve()
        with catalog_path.open("r", encoding="utf-8") as fp:
            data = json.load(fp)
        root = catalog_path.parent
        templates: dict[str, ResolvedTemplate] = {}
        for template_id, item in data.get("templates", {}).items():
            templates[template_id] = ResolvedTemplate(
                template_path=(root / item["template_path"]).resolve(),
                style_profile_path=(
                    (root / item["style_profile_path"]).resolve()
                    if item.get("style_profile_path")
                    else None
                ),
                anchor=item.get("anchor", DEFAULT_TEMPLATE_ANCHOR),
            )
        return cls(root=root, templates=templates)

    def resolve(self, template_id: str) -> ResolvedTemplate:
        try:
            return self.templates[template_id]
        except KeyError as exc:
            raise KeyError(f"Unknown template id: {template_id!r}") from exc


def create_starter_template(
    output_path: str | Path,
    *,
    title: str = "PaddleOCR Markdown Template",
    anchor: str = DEFAULT_TEMPLATE_ANCHOR,
) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(anchor)
    doc.save(output)
    return output


def render_markdown_to_docx(
    *,
    markdown_text: str,
    output_path: str | Path,
    template_path: str | Path | None = None,
    anchor: str = DEFAULT_TEMPLATE_ANCHOR,
    style_profile: StyleProfile | None = None,
) -> Path:
    profile = style_profile or StyleProfile()
    doc = Document(template_path) if template_path else Document()
    _apply_document_defaults(doc, profile)

    html = markdown.markdown(
        markdown_text,
        extensions=["tables", "fenced_code", "sane_lists"],
    )
    soup = BeautifulSoup(html, "html.parser")
    renderer = _MarkdownDocxRenderer(doc=doc, profile=profile)

    if template_path is None:
        renderer.append_nodes(soup.contents)
    else:
        anchor_paragraph = _find_anchor_paragraph(doc, anchor)
        renderer.insert_after_anchor(anchor_paragraph, soup.contents)
        _delete_paragraph(anchor_paragraph)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


class _MarkdownDocxRenderer:
    def __init__(self, *, doc: Document, profile: StyleProfile):
        self.doc = doc
        self.profile = profile
        self._table_parser = HtmlToDocx()

    def append_nodes(self, nodes: Iterable[Any]) -> None:
        for node in nodes:
            self._append_node(node)

    def insert_after_anchor(self, paragraph, nodes: Iterable[Any]) -> None:
        anchor_element = paragraph._p
        for node in nodes:
            anchor_element = self._insert_node_after(node, anchor_element)

    def _append_node(self, node: Any) -> None:
        if isinstance(node, NavigableString):
            if node.strip():
                paragraph = self.doc.add_paragraph(node.strip())
                _apply_block_style(paragraph, self.profile.normal)
            return
        if not isinstance(node, Tag):
            return
        tag_name = node.name.lower()
        if tag_name in {"h1", "h2", "h3"}:
            paragraph = self.doc.add_paragraph(node.get_text(" ", strip=True))
            style = {
                "h1": self.profile.heading1,
                "h2": self.profile.heading2,
                "h3": self.profile.heading3,
            }[tag_name]
            _apply_block_style(paragraph, style)
        elif tag_name == "p":
            self._append_paragraph_like(node, self.profile.normal)
        elif tag_name in {"ul", "ol"}:
            self._append_list(node, ordered=(tag_name == "ol"))
        elif tag_name == "pre":
            paragraph = self.doc.add_paragraph(node.get_text("\n", strip=True))
            _apply_block_style(paragraph, self.profile.code_block)
        elif tag_name == "blockquote":
            self._append_paragraph_like(node, self.profile.quote)
        elif tag_name == "table":
            self._append_table(str(node))
        elif tag_name == "div":
            self._append_div(node)
        elif tag_name == "img":
            self._append_image(node)

    def _insert_node_after(self, node: Any, anchor_element):
        if isinstance(node, NavigableString):
            if not node.strip():
                return anchor_element
            paragraph = self.doc.add_paragraph(node.strip())
            _move_element_after(paragraph._p, anchor_element)
            _apply_block_style(paragraph, self.profile.normal)
            return paragraph._p
        if not isinstance(node, Tag):
            return anchor_element
        tag_name = node.name.lower()
        if tag_name in {"h1", "h2", "h3"}:
            paragraph = self.doc.add_paragraph(node.get_text(" ", strip=True))
            _move_element_after(paragraph._p, anchor_element)
            style = {
                "h1": self.profile.heading1,
                "h2": self.profile.heading2,
                "h3": self.profile.heading3,
            }[tag_name]
            _apply_block_style(paragraph, style)
            return paragraph._p
        if tag_name == "p":
            paragraph = self.doc.add_paragraph()
            _move_element_after(paragraph._p, anchor_element)
            self._fill_paragraph(paragraph, node, self.profile.normal)
            return paragraph._p
        if tag_name in {"ul", "ol"}:
            current = anchor_element
            for li in node.find_all("li", recursive=False):
                paragraph = self.doc.add_paragraph(style="List Number" if tag_name == "ol" else "List Bullet")
                _move_element_after(paragraph._p, current)
                self._fill_paragraph(paragraph, li, self.profile.list_item)
                current = paragraph._p
            return current
        if tag_name == "pre":
            paragraph = self.doc.add_paragraph(node.get_text("\n", strip=True))
            _move_element_after(paragraph._p, anchor_element)
            _apply_block_style(paragraph, self.profile.code_block)
            return paragraph._p
        if tag_name == "blockquote":
            paragraph = self.doc.add_paragraph()
            _move_element_after(paragraph._p, anchor_element)
            self._fill_paragraph(paragraph, node, self.profile.quote)
            return paragraph._p
        if tag_name == "table":
            table = self._append_table(str(node))
            _move_element_after(table._tbl, anchor_element)
            return table._tbl
        if tag_name == "div":
            return self._insert_div_after(node, anchor_element)
        if tag_name == "img":
            paragraph = self._append_image(node)
            if paragraph is None:
                return anchor_element
            _move_element_after(paragraph._p, anchor_element)
            return paragraph._p
        return anchor_element

    def _append_paragraph_like(self, node: Tag, style: BlockStyle) -> None:
        paragraph = self.doc.add_paragraph()
        self._fill_paragraph(paragraph, node, style)

    def _fill_paragraph(self, paragraph, node: Tag, style: BlockStyle) -> None:
        text = node.get_text(" ", strip=True)
        if not text and node.find("img"):
            img_paragraph = self._append_image(node.find("img"))
            if img_paragraph is not None:
                paragraph._element.getparent().remove(paragraph._element)
            return
        run = paragraph.add_run(text)
        _apply_run_style(run, style)
        _apply_paragraph_style(paragraph, style)

    def _append_list(self, node: Tag, *, ordered: bool) -> None:
        for li in node.find_all("li", recursive=False):
            paragraph = self.doc.add_paragraph(style="List Number" if ordered else "List Bullet")
            self._fill_paragraph(paragraph, li, self.profile.list_item)

    def _append_table(self, table_html: str):
        existing_tables = len(self.doc.tables)
        self._table_parser.handle_table(table_html, self.doc)
        table = self.doc.tables[existing_tables]
        _apply_table_style(table, self.profile)
        return table

    def _append_div(self, node: Tag) -> None:
        image = node.find("img")
        if image is not None:
            self._append_image(image, centered=node.get("align") == "center")
            return
        paragraph = self.doc.add_paragraph(node.get_text(" ", strip=True))
        _apply_block_style(paragraph, self.profile.normal)

    def _insert_div_after(self, node: Tag, anchor_element):
        image = node.find("img")
        if image is not None:
            paragraph = self._append_image(image, centered=node.get("align") == "center")
            if paragraph is None:
                return anchor_element
            _move_element_after(paragraph._p, anchor_element)
            return paragraph._p
        paragraph = self.doc.add_paragraph(node.get_text(" ", strip=True))
        _move_element_after(paragraph._p, anchor_element)
        _apply_block_style(paragraph, self.profile.normal)
        return paragraph._p

    def _append_image(self, node: Tag | None, *, centered: bool = True):
        if node is None:
            return None
        src = node.get("src")
        if not src:
            return None
        image_path = Path(src)
        if not image_path.exists():
            paragraph = self.doc.add_paragraph(f"[missing image: {src}]")
            _apply_block_style(paragraph, self.profile.quote)
            return paragraph
        paragraph = self.doc.add_paragraph()
        if centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(5.5))
        return paragraph


def _find_anchor_paragraph(doc: Document, anchor: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == anchor:
            return paragraph
    raise ValueError(f"Unable to find template anchor {anchor!r} in the template document.")


def _apply_document_defaults(doc: Document, profile: StyleProfile) -> None:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = profile.normal.font_name
    normal_style.font.size = Pt(profile.normal.font_size_pt)
    if profile.normal.east_asia_font:
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), profile.normal.east_asia_font)


def _apply_block_style(paragraph, style: BlockStyle) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        _apply_run_style(run, style)
    _apply_paragraph_style(paragraph, style)


def _apply_run_style(run, style: BlockStyle) -> None:
    run.font.name = style.font_name
    run.font.size = Pt(style.font_size_pt)
    run.font.bold = style.bold
    run.font.italic = style.italic
    run.font.underline = style.underline
    if style.east_asia_font:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), style.east_asia_font)


def _apply_paragraph_style(paragraph, style: BlockStyle) -> None:
    alignment = style.alignment.lower()
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(style.space_before_pt)
    paragraph_format.space_after = Pt(style.space_after_pt)
    if style.line_spacing is not None:
        paragraph_format.line_spacing = style.line_spacing


def _apply_table_style(table, profile: StyleProfile) -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _apply_block_style(
                    paragraph,
                    profile.table_header if row_idx == 0 else profile.table_cell,
                )


def _move_element_after(element, anchor_element) -> None:
    anchor_element.addnext(element)


def _delete_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None

